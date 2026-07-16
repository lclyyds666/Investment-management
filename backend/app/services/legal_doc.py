"""生成「法律文件审批表」.docx(严格 3cm 行高 + 方正小标宋简体/仿宋_GB2312)。

参考模板:`供管法律文件审批表.doc`。用 python-docx(已装,无新依赖)精确控制
行高与字体,用户用 Word 打开打印,字体在本机呈现(出版单位通常已装公文字体)。
"""
from __future__ import annotations

import base64
import io
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

TITLE_TEXT = "山东出版供应链法律文件审批表"
PUBLISHER = "出版供应链"          # 发文单位默认值
TITLE_FONT = "方正小标宋简体"      # 标题字体(公文标准)
BODY_FONT = "仿宋_GB2312"          # 正文字体(公文标准)
ROW_HEIGHT_CM = 3                  # 严格行高 3cm

# 4 个意见栏(角色值 → 栏目名),顺序即审批流转顺序
OPINION_ROLES = [
    ("scm_director", "供管公司负责人意见"),
    ("legal_counsel", "法律顾问意见"),
    ("risk_auditor", "投资公司法务风控部意见"),
    ("invest_director", "投资公司分管领导意见"),
]


def _set_font(run, name: str, size_pt: float, bold: bool = False) -> None:
    """设置 run 字体(含中文 eastAsia,确保 Word 用指定中文字体渲染)。"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


# docx 可内嵌的光栅图片格式（SVG 等矢量格式不支持 → 回退姓名文本）
_RASTER_SUBTYPES = {"png", "jpeg", "jpg", "gif", "bmp"}


def _parse_signature(uri: str):
    """解析电子签名 data-URI(data:image/*;base64,XXXX) → 原始字节；

    非法 / 空 / 非光栅格式(如 svg) 一律返回 None，交由调用方回退为姓名占位。
    """
    if not uri or not uri.startswith("data:image/"):
        return None
    m = re.match(r"data:image/([a-zA-Z0-9.+-]+);base64,(.*)$", uri, re.DOTALL)
    if not m:
        return None
    if m.group(1).lower() not in _RASTER_SUBTYPES:
        return None
    try:
        return base64.b64decode(m.group(2))
    except Exception:  # noqa: BLE001 解码失败即回退姓名
        return None


def _fill_opinion_cell(cell, op: dict | None) -> None:
    """渲染单个审批意见栏：实际审批意见 + 电子签名(图片占位/姓名) + 日期。"""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in list(p0.runs):
        r.text = ""
    if not op:
        return  # 该环节尚无审批记录 → 留空

    # 1) 审批意见：读取实际 comment（可能为空，不再默认“同意”）
    comment = str(op.get("comment") or "")
    first = True
    for line in (comment.split("\n") if comment else [""]):
        para = p0 if first else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_font(para.add_run(line), BODY_FONT, 14)
        first = False

    # 2) 签名 + 日期（右对齐）：已上传签名渲染图片，否则姓名占位
    sig_para = cell.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(sig_para.add_run("签名："), BODY_FONT, 12)
    raw = _parse_signature(op.get("signature") or "")
    name = str(op.get("approver_name") or "")
    embedded = False
    if raw:
        try:
            sig_para.add_run().add_picture(io.BytesIO(raw), height=Cm(1.1))
            embedded = True
        except Exception:  # noqa: BLE001 图片异常回退姓名
            embedded = False
    if not embedded:
        _set_font(sig_para.add_run(name), BODY_FONT, 14)  # 未上传签名 → 姓名占位
    date = str(op.get("date") or "")
    if date:
        _set_font(sig_para.add_run("　" + date), BODY_FONT, 12)


def _fill_cell(cell, text: str, *, font: str = BODY_FONT, size: float = 14,
               bold: bool = False, center: bool = False) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lines = str(text or "").split("\n")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(lines[0])
    _set_font(run, font, size, bold)
    for extra in lines[1:]:
        np = cell.add_paragraph()
        np.alignment = p.alignment
        _set_font(np.add_run(extra), font, size, bold)


def build_legal_doc(contract: dict, opinions: dict) -> bytes:
    """生成审批表 .docx 字节。

    contract: {title, contract_no, sign_date, creator_name}
    opinions: {role_value: {comment, approver_name, signature, date}}
             comment 为实际审批意见（空则留空，不默认“同意”）；
             signature 为电子签名 data-URI（光栅图渲染图片，否则用姓名占位）。
    """
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2)
    sec.left_margin = sec.right_margin = Cm(2.5)

    # 标题
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(h.add_run(TITLE_TEXT), TITLE_FONT, 22, bold=True)
    doc.add_paragraph()  # 标题与表格间距

    table = doc.add_table(rows=0, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def add_row():
        row = table.add_row()
        row.height = Cm(ROW_HEIGHT_CM)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        return row

    # 第一行:发文单位 / 发文时间 / 发文人员
    c = add_row().cells
    _fill_cell(c[0], "发文单位", bold=True, center=True)
    _fill_cell(c[1], PUBLISHER, center=True)
    _fill_cell(c[2], "发文时间", bold=True, center=True)
    _fill_cell(c[3], contract.get("sign_date") or "", center=True)
    _fill_cell(c[4], "发文人员", bold=True, center=True)
    _fill_cell(c[5], contract.get("creator_name") or "", center=True)

    # 文件名称（合同编号）—— 合并为一行，格式「文件名称（合同编号）」
    r = add_row()
    _fill_cell(r.cells[0], "文件名称", bold=True, center=True)
    _title = contract.get("title") or ""
    _no = contract.get("contract_no") or ""
    _name_value = f"{_title}（{_no}）" if _no else _title
    _fill_cell(r.cells[1].merge(r.cells[5]), _name_value)

    # 4 个意见栏：实际审批意见 + 电子签名占位
    for role, label in OPINION_ROLES:
        r = add_row()
        _fill_cell(r.cells[0], label, bold=True, center=True)
        _fill_opinion_cell(r.cells[1].merge(r.cells[5]), opinions.get(role))

    # 标签列固定宽
    for row in table.rows:
        row.cells[0].width = Cm(3.2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
