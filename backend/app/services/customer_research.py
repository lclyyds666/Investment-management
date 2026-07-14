"""客户 AI 尽职调查服务：准入资料解析 + 联网检索(博查) + DeepSeek 综合。

链路：
  上传资料 → 逐页文本(保留页码) → 正则抽取关键信息
                                   ↘
  外部：博查 Web 搜索(带 URL 的最新资讯) ↘
                                   都作为 Context 喂给 DeepSeek
  → 产出四段式《尽职调查报告》(基础概况/经营分析/外部舆情/AI风险预警) + 合作建议 + 来源标注

降级哲学（与项目一致，接口永不 500）：
  - 无 DEEPSEEK_API_KEY / 调用失败 → 规则引擎兜底成文；
  - 无 SEARCH_API_KEY / 搜索失败 → 跳过联网，"外部舆情"段标注"未联网核实"。
"""
from __future__ import annotations

import io
import json
import logging
import re
import uuid
from pathlib import Path
from typing import Any

import httpx

from app.core.config import settings

logger = logging.getLogger("app.customer_research")

RECOMMENDATIONS = ("优先合作", "谨慎合作", "严禁合作")
_MAX_TEXT_FOR_LLM = 6000  # 喂给大模型的内部资料截断长度（控制 token）


# --------------------------------------------------------------------------- #
# 一、文件解析（PDF / Docx → 逐页文本）
# --------------------------------------------------------------------------- #
def extract_pages(filename: str, content: bytes) -> tuple[str, list[dict]]:
    """返回 (file_type, pages)；pages=[{'page':1,'text':...}]。仅支持 pdf/docx。"""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return "pdf", _extract_pdf(content)
    if name.endswith(".docx"):
        return "docx", _extract_docx(content)
    raise ValueError("仅支持 .pdf / .docx 格式（旧 .doc、扫描件图片型 PDF 暂不支持）")


def _extract_pdf(content: bytes) -> list[dict]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001 个别页解析异常不影响整体
            text = ""
        pages.append({"page": i, "text": text.strip()})
    return pages


def _extract_docx(content: bytes) -> list[dict]:
    import docx

    doc = docx.Document(io.BytesIO(content))
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # 表格文本也纳入
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    # docx 无严格分页，整篇作为第 1 页
    return [{"page": 1, "text": "\n".join(paras).strip()}]


# --------------------------------------------------------------------------- #
# 二、关键信息抽取（正则，离线零依赖）
# --------------------------------------------------------------------------- #
def extract_key_info(full_text: str) -> dict:
    def find(*patterns: str) -> str:
        for pat in patterns:
            m = re.search(pat, full_text)
            if m:
                return m.group(1).strip()
        return ""

    return {
        "credit_code": find(r"(?:统一社会信用代码|信用代码)[：:\s]*([0-9A-Z]{18})"),
        "capital": find(r"(?:注册资本|注册资金)[：:\s]*([^\n，,。;；]{1,30})"),
        "legal_person": find(r"(?:法定代表人|法人代表|法人)[：:\s]*([^\n，,。;；]{1,20})"),
        "established": find(r"(?:成立日期|成立时间|注册日期)[：:\s]*([0-9\-/年月日]{4,12})"),
        "scope": find(r"(?:经营范围|业务范围)[：:\s]*([^\n]{1,200})"),
    }


# --------------------------------------------------------------------------- #
# 三、原始文件落盘
# --------------------------------------------------------------------------- #
def save_file(customer_id: int, filename: str, content: bytes) -> str:
    """把原始文件存到 UPLOAD_DIR/customer_{id}/ 下，返回磁盘存储名(uuid+扩展名)。"""
    ext = Path(filename).suffix.lower()
    stored = f"{uuid.uuid4().hex}{ext}"
    d = Path(settings.UPLOAD_DIR) / f"customer_{customer_id}"
    d.mkdir(parents=True, exist_ok=True)
    (d / stored).write_bytes(content)
    return stored


def file_path(customer_id: int, stored_name: str) -> Path:
    return Path(settings.UPLOAD_DIR) / f"customer_{customer_id}" / stored_name


# --------------------------------------------------------------------------- #
# 四、联网检索（博查 Bocha Web Search）
# --------------------------------------------------------------------------- #
def web_search(query: str) -> list[dict]:
    """调用博查搜索公司资讯；未配置 key 或失败时返回 []（触发降级）。"""
    if not settings.SEARCH_ENABLED:
        return []
    try:
        resp = httpx.post(
            settings.SEARCH_BASE_URL,
            headers={
                "Authorization": f"Bearer {settings.SEARCH_API_KEY}",
                "Content-Type": "application/json",
            },
            json={"query": query, "summary": True, "count": settings.SEARCH_COUNT, "page": 1},
            timeout=settings.AI_TIMEOUT_SECONDS,
        )
        resp.raise_for_status()
        data = resp.json()
        # 博查响应：data.data.webPages.value[]
        values = (((data or {}).get("data") or {}).get("webPages") or {}).get("value") or []
        results = []
        for p in values:
            item = {
                "title": (p.get("name") or "").strip(),
                "url": (p.get("url") or "").strip(),
                "snippet": (p.get("summary") or p.get("snippet") or "").strip(),
                "site": (p.get("siteName") or "").strip(),
                "date": (p.get("datePublished") or p.get("dateLastCrawled") or "").strip()[:10],
            }
            if item["title"] and item["url"]:
                results.append(item)
        return results
    except Exception as exc:  # noqa: BLE001 搜索失败不阻断报告生成
        logger.warning("博查搜索失败，降级为无外部资讯：%s", exc)
        return []


# --------------------------------------------------------------------------- #
# 五、生成《尽职调查报告》
# --------------------------------------------------------------------------- #
def generate_report(
    customer_name: str,
    full_text: str,
    key_info: dict,
    web_results: list[dict],
    materials_meta: list[dict],
) -> dict:
    """融合内部资料 + 外部资讯，优先 DeepSeek 综合，失败/未配置回退规则引擎。

    返回 {sections, recommendation, report_md, sources, engine, searched, search_count}。
    """
    # 统一构造来源清单：外部网页（编号）+ 内部资料（文件+页数）
    sources: list[dict] = []
    for i, w in enumerate(web_results, start=1):
        sources.append({"no": i, "type": "web", "title": w["title"], "ref": w["url"],
                        "site": w.get("site", ""), "date": w.get("date", "")})
    for m in materials_meta:
        sources.append({"type": "doc", "title": m["filename"], "ref": f"共 {m['page_count']} 页"})

    searched = len(web_results) > 0

    if settings.AI_ENABLED:
        try:
            report = _report_by_llm(customer_name, full_text, key_info, web_results)
            report["engine"] = "deepseek"
        except Exception as exc:  # noqa: BLE001
            logger.warning("DeepSeek 生成尽调报告失败，回退规则引擎：%s", exc)
            report = _report_by_rules(customer_name, key_info, web_results)
            report["engine"] = "rule"
    else:
        report = _report_by_rules(customer_name, key_info, web_results)
        report["engine"] = "rule"

    report["sources"] = sources
    report["searched"] = searched
    report["search_count"] = len(web_results)
    report["report_md"] = _render_markdown(customer_name, report, sources)
    return report


_SYSTEM_PROMPT = (
    "你是一名资深企业尽职调查(尽调)分析师，服务于山东出版供应链管理有限公司的客户准入风控。"
    "你将收到某潜在客户的【内部资料摘录】与【外部网络资讯】，请严格依据这些材料撰写一份专业、"
    "客观、可执行的《客户尽职调查报告》。要求：\n"
    "1. **以【内部资料】为核心依据做深度解析**（公司主体、注册资本、经营范围、核心业务），"
    "结合其业务模式与所处行业发展趋势展开分析；\n"
    "2. **严格甄别外部资讯的相关性**：网络资讯常混入与本公司无关的同名企业、股吧/论坛个人言论、"
    "广告招聘、其他行业公司等噪声——**必须剔除这些无关信息，只采用与目标公司确切相关的内容**，"
    "无关项不得写入报告、不得作为结论依据；若外部资讯整体与目标公司无关，请在【外部舆情】如实说明"
    "'未检索到与该公司确切相关的外部资讯'；\n"
    "3. 只依据给定材料推断，不臆造不存在的具体数字或事实；材料不足处如实说明'资料未提供'；\n"
    "4. 涉及外部资讯的结论，请在句末用 [编号] 标注所引用且确认相关的来源；\n"
    "5. 合作建议必须是且仅是三选一：优先合作 / 谨慎合作 / 严禁合作；\n"
    "6. 严格输出 JSON，不要输出任何多余文字或 Markdown 代码块。"
)

_JSON_SPEC = (
    "请仅输出如下 JSON：\n"
    "{\n"
    '  "basic": "【基础概况】基于内部资料总结的公司主体、注册资本、经营范围、核心业务等(120-200字)",\n'
    '  "operation": "【经营分析】结合业务模式与行业趋势分析业务可行性(120-200字)",\n'
    '  "sentiment": "【外部舆情】基于外部资讯评估信用风险/合规性/诉讼舆情，句末标注[编号](120-200字)",\n'
    '  "risk": "【AI风险预警】综合研判，给出关键风险点(100-180字)",\n'
    '  "recommendation": "优先合作|谨慎合作|严禁合作",\n'
    '  "key_risks": ["关键风险点1", "关键风险点2"]\n'
    "}"
)


def _report_by_llm(customer_name: str, full_text: str, key_info: dict, web_results: list[dict]) -> dict:
    from openai import OpenAI

    client = OpenAI(
        api_key=settings.DEEPSEEK_API_KEY,
        base_url=settings.DEEPSEEK_BASE_URL,
        timeout=settings.AI_TIMEOUT_SECONDS,
    )

    internal = full_text[:_MAX_TEXT_FOR_LLM]
    if len(full_text) > _MAX_TEXT_FOR_LLM:
        internal += "\n……(资料过长已截断)"

    ki = "、".join(f"{k}={v}" for k, v in key_info.items() if v) or "（正则未抽取到结构化字段）"
    if web_results:
        ext_lines = "\n".join(
            f"[{i}] {w['title']}（{w.get('site','')} {w.get('date','')}）{w['url']}\n    摘要：{w['snippet']}"
            for i, w in enumerate(web_results, start=1)
        )
    else:
        ext_lines = "（未配置联网搜索或未检索到外部资讯，请在【外部舆情】段注明'未联网核实，以下基于模型既有认知'）"

    user_prompt = (
        f"目标公司名称：{customer_name}\n\n"
        f"【关键信息(正则抽取)】\n{ki}\n\n"
        f"【内部资料摘录】\n{internal or '（未提取到文本，可能为扫描件）'}\n\n"
        f"【外部网络资讯】（可能含无关噪声，请仅采用与「{customer_name}」确切相关者，无关项忽略）\n{ext_lines}\n\n"
        + _JSON_SPEC
    )

    resp = client.chat.completions.create(
        model=settings.DEEPSEEK_MODEL,
        messages=[
            {"role": "system", "content": _SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        response_format={"type": "json_object"},
        temperature=0.4,
        stream=False,
    )
    data = json.loads(resp.choices[0].message.content or "{}")

    rec = str(data.get("recommendation", "")).strip()
    if rec not in RECOMMENDATIONS:
        rec = "谨慎合作"
    sections = {
        "basic": str(data.get("basic", "")).strip(),
        "operation": str(data.get("operation", "")).strip(),
        "sentiment": str(data.get("sentiment", "")).strip(),
        "risk": str(data.get("risk", "")).strip(),
        "key_risks": [str(x).strip() for x in (data.get("key_risks") or []) if str(x).strip()],
    }
    # 任一段为空则用规则兜底补齐，避免前端空白
    if not all([sections["basic"], sections["operation"], sections["sentiment"], sections["risk"]]):
        fb = _report_by_rules(customer_name, key_info, web_results)
        for k in ("basic", "operation", "sentiment", "risk"):
            sections[k] = sections[k] or fb["sections"][k]
        sections["key_risks"] = sections["key_risks"] or fb["sections"]["key_risks"]
        rec = rec or fb["recommendation"]
    return {"sections": sections, "recommendation": rec}


def _report_by_rules(customer_name: str, key_info: dict, web_results: list[dict]) -> dict:
    """无大模型时的规则兜底：把已抽取信息结构化成文，保证四段有内容。"""
    cap = key_info.get("capital") or "资料未提供"
    scope = key_info.get("scope") or "资料未提供"
    legal = key_info.get("legal_person") or "资料未提供"
    code = key_info.get("credit_code") or "资料未提供"

    basic = (
        f"{customer_name}：统一社会信用代码 {code}，法定代表人 {legal}，注册资本 {cap}。"
        f"经营范围：{scope}。以上为内部准入资料所载核心信息。"
    )
    operation = (
        "基于经营范围与业务模式判断，需结合文旅供应链行业景气度与公司实际履约能力评估业务可行性；"
        "建议进一步核验其主营业务的真实交易记录与上下游稳定性。（未接入大模型，为规则化摘要）"
    )
    if web_results:
        cites = "；".join(f"{w['title']}[{i}]" for i, w in enumerate(web_results, start=1))
        sentiment = f"检索到 {len(web_results)} 条外部资讯，涉及：{cites}。请人工复核其中的诉讼/信用与舆情信息。"
    else:
        sentiment = "未联网核实（未配置搜索 API 或未检索到结果）。建议接入企业征信/裁判文书数据源后再评估外部合规风险。"
    risk = (
        "关键风险点：①资料完整性需人工复核；②外部信用/诉讼信息尚未权威核验。"
        "在补充核验前，建议以稳健口径推进合作。"
    )
    return {
        "sections": {
            "basic": basic, "operation": operation, "sentiment": sentiment, "risk": risk,
            "key_risks": ["资料完整性待复核", "外部信用/诉讼待权威核验"],
        },
        "recommendation": "谨慎合作",
    }


def _render_markdown(customer_name: str, report: dict, sources: list[dict]) -> str:
    s = report["sections"]
    lines = [
        f"# 《{customer_name}》客户尽职调查报告",
        f"\n**合作建议：{report['recommendation']}**\n",
        "## 一、基础概况", s["basic"],
        "\n## 二、经营分析", s["operation"],
        "\n## 三、外部舆情", s["sentiment"],
        "\n## 四、AI 风险预警", s["risk"],
    ]
    if s.get("key_risks"):
        lines.append("\n**关键风险点：**")
        lines += [f"- {r}" for r in s["key_risks"]]
    lines.append("\n## 信息来源")
    web = [x for x in sources if x["type"] == "web"]
    doc = [x for x in sources if x["type"] == "doc"]
    if web:
        lines += [f"[{x['no']}] {x['title']}（{x.get('site','')} {x.get('date','')}）{x['ref']}" for x in web]
    else:
        lines.append("（外部资讯：未联网核实）")
    for x in doc:
        lines.append(f"- 内部资料：{x['title']}（{x['ref']}）")
    engine_note = "DeepSeek 大模型综合" if report.get("engine") == "deepseek" else "规则引擎(未接大模型)"
    lines.append(f"\n> 生成引擎：{engine_note}；外部资讯 {report.get('search_count', 0)} 条。")
    return "\n".join(lines)
